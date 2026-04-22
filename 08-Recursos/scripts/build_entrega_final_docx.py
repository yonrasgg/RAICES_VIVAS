#!/usr/bin/env python3
"""Generate 'Entrega Final Proyecto Raices Vivas.docx' from the Markdown master,
using the official CENFOTEC template as base.

Resolves Obsidian-specific syntax so the DOCX is self-contained:
- ![[File#Section]]  transclusions   → inlines target section (recursively)
- ```mermaid blocks                  → renders PNG via npx mmdc (cached)
- ```sqlseal blocks                  → executes against an in-memory SQLite
                                       built from all MD frontmatter and
                                       renders the result as a DOCX table
- [[wiki|alias]] / [[wiki]]          → plain text
"""

from __future__ import annotations

import copy
import hashlib
import os
import re
import shutil
import sqlite3
import subprocess
import tempfile
from pathlib import Path

import yaml
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent.parent
TEMPLATE = ROOT / "08-Recursos" / "PDFs" / "01-Formato Entrega Final.docx"
SRC_MD = ROOT / "06-Entregables" / "Entrega-Final" / "Raíces Vivas — Entrega Final Implementación y Gestión.md"
OUT_DIR = ROOT / "06-Entregables" / "Entrega-Final"
OUT = OUT_DIR / "Entrega Final Proyecto Raices Vivas.docx"
MERMAID_CACHE = ROOT / "08-Recursos" / "scripts" / ".mermaid_cache"
ACTOR_DIAGRAM_PATH = ROOT / "08-Recursos" / "Imágenes" / "diagram_actor_modulo.png"
USECASE_DIAGRAM_PATH = ROOT / "08-Recursos" / "Imágenes" / "diagram_usecase.png"

MERMAID_CACHE.mkdir(parents=True, exist_ok=True)

COVER_REPLACEMENTS = {
    "<Título del Trabajo>": "Raíces Vivas — Entrega Final: Implementación y Gestión",
    "<Autor(es)>": "Autores",
    "<Nombre Estudiante 1>": "Geovanny Alpízar Sandino",
    "<Nombre Estudiante 2>": "Elkin Cerda González",
    "<Nombre del Curso>": "Introducción a la Ingeniería del Software (SOFT-09, SCV7)",
    "<Nombre del Profesor a Cargo>": "Johnny Marín Sánchez",
    "<Mes, Año>": "Abril, 2026",
}

# ═══════════════════════════════════════════════════════════════════════════
# 1. Vault indexing + SQLSeal emulator
# ═══════════════════════════════════════════════════════════════════════════

IGNORE_PREFIXES = (
    ".obsidian", ".git", ".venv", "node_modules", "app/node_modules",
    "app/dist", "09-QA",
)

_MD_CACHE: dict[str, str] = {}


def read_md(path: Path) -> str:
    key = str(path)
    if key not in _MD_CACHE:
        _MD_CACHE[key] = path.read_text(encoding="utf-8", errors="ignore")
    return _MD_CACHE[key]


def parse_frontmatter(text: str) -> dict | None:
    if not text.startswith("---"):
        return None
    end = text.find("\n---", 3)
    if end == -1:
        return None
    try:
        data = yaml.safe_load(text[3:end]) or {}
    except yaml.YAMLError:
        return None
    return data if isinstance(data, dict) else None


def build_vault_db() -> sqlite3.Connection:
    """Scan vault for MD files, parse YAML frontmatter, populate SQLite `files`."""
    records: list[dict] = []
    for md in ROOT.rglob("*.md"):
        rel = str(md.relative_to(ROOT))
        if any(rel.startswith(p) for p in IGNORE_PREFIXES):
            continue
        text = read_md(md)
        fm = parse_frontmatter(text)
        if fm is None:
            continue
        fm["path"] = rel
        fm["name"] = md.stem
        records.append(fm)

    keys: set[str] = set()
    for r in records:
        keys.update(r.keys())
    keys.update({"path", "name", "type", "title", "module", "wbs", "priority",
                 "status", "category", "date", "assignee", "effort_actual",
                 "severity", "strategy", "sprint"})
    cols = sorted(keys)

    conn = sqlite3.connect(":memory:")
    quoted = ", ".join(f'"{c}" TEXT' for c in cols)
    conn.execute(f"CREATE TABLE files ({quoted})")
    placeholders = ",".join("?" for _ in cols)
    quoted_cols = ", ".join(f'"{c}"' for c in cols)
    for r in records:
        vals = []
        for c in cols:
            v = r.get(c)
            if v is None:
                vals.append(None)
            elif isinstance(v, (list, dict)):
                vals.append(str(v))
            else:
                vals.append(str(v))
        conn.execute(
            f"INSERT INTO files ({quoted_cols}) VALUES ({placeholders})", vals
        )
    conn.commit()
    print(f"  Vault DB: {len(records)} files indexed, {len(cols)} columns")
    return conn


def run_sqlseal(conn: sqlite3.Connection, sql: str):
    sql = sql.strip().rstrip(";")
    cur = conn.execute(sql)
    headers = [d[0] for d in cur.description]
    rows = [[("" if v is None else str(v)) for v in row] for row in cur.fetchall()]
    return headers, rows


# ═══════════════════════════════════════════════════════════════════════════
# 2. Transclusion resolver  ![[File#Section]]
# ═══════════════════════════════════════════════════════════════════════════

_FILE_INDEX: dict[str, Path] | None = None


def get_file_index() -> dict[str, Path]:
    global _FILE_INDEX
    if _FILE_INDEX is None:
        idx: dict[str, Path] = {}
        for md in ROOT.rglob("*.md"):
            rel = str(md.relative_to(ROOT))
            if any(rel.startswith(p) for p in IGNORE_PREFIXES):
                continue
            if md.stem not in idx or len(str(md)) < len(str(idx[md.stem])):
                idx[md.stem] = md
        _FILE_INDEX = idx
    return _FILE_INDEX


def _norm_header(s: str) -> str:
    return re.sub(r"\s+", " ", s).strip()


def extract_section(text: str, section_name: str) -> str | None:
    lines = text.split("\n")
    target = _norm_header(section_name)
    section_re = re.compile(r"^(#+)\s+(.+?)\s*$")
    start: int | None = None
    level = 0
    for i, line in enumerate(lines):
        m = section_re.match(line)
        if not m:
            continue
        if start is None:
            if _norm_header(m.group(2)) == target:
                start = i + 1
                level = len(m.group(1))
        else:
            if len(m.group(1)) <= level:
                return "\n".join(lines[start:i]).strip()
    if start is not None:
        return "\n".join(lines[start:]).strip()
    return None


TRANSCLUSION_RE = re.compile(r"!\[\[([^\]]+)\]\]")


def resolve_transcludes(text: str, depth: int = 0) -> str:
    if depth > 4:
        return text

    def repl(m: re.Match) -> str:
        ref = m.group(1)
        if "#" in ref:
            file_part, header = ref.split("#", 1)
        else:
            file_part, header = ref, None
        file_part = file_part.split("|", 1)[0].strip()
        idx = get_file_index()
        target = idx.get(file_part)
        if target is None:
            for k, v in idx.items():
                if k.lower() == file_part.lower():
                    target = v; break
        if target is None:
            return f"*[Referencia no encontrada: {ref}]*"
        target_text = read_md(target)
        target_text = re.sub(r"^---\n.*?\n---\n", "", target_text, count=1, flags=re.DOTALL)
        if header:
            section = extract_section(target_text, header)
            if section is None:
                return f"*[Sección no encontrada: {ref}]*"
            content = section
        else:
            content = target_text
        return resolve_transcludes(content, depth + 1)

    return TRANSCLUSION_RE.sub(repl, text)


# ═══════════════════════════════════════════════════════════════════════════
# 3. Mermaid rendering via npx mmdc (cached)
# ═══════════════════════════════════════════════════════════════════════════


def render_mermaid_to_png(code: str) -> Path | None:
    if "Actores Primarios" in code and "Módulos del Sistema" in code:
        if ACTOR_DIAGRAM_PATH.exists():
            return ACTOR_DIAGRAM_PATH
    if "CU_EDU_01" in code and "CU_SAB_01" in code and "CU_SAL_01" in code:
        if USECASE_DIAGRAM_PATH.exists():
            return USECASE_DIAGRAM_PATH

    digest = hashlib.sha256(code.encode("utf-8")).hexdigest()[:12]
    png_file = MERMAID_CACHE / f"mmdc_{digest}.png"
    if png_file.exists():
        return png_file

    with tempfile.NamedTemporaryFile(mode="w", suffix=".mmd", delete=False, encoding="utf-8") as tmp:
        tmp.write(code)
        tmp_path = tmp.name
    try:
        result = subprocess.run(
            ["npx", "mmdc", "-i", tmp_path, "-o", str(png_file),
             "-b", "white", "-t", "default", "--scale", "3"],
            capture_output=True, text=True, timeout=180, cwd=str(ROOT),
        )
        if result.returncode != 0:
            print(f"  ⚠ mmdc failed: {result.stderr[:250]}")
            return None
    finally:
        os.unlink(tmp_path)
    return png_file if png_file.exists() else None


# ═══════════════════════════════════════════════════════════════════════════
# 4. Markdown pre-processing
# ═══════════════════════════════════════════════════════════════════════════

FRONTMATTER_RE = re.compile(r"^---\n.*?\n---\n", re.DOTALL)
WIKILINK_ALIAS_RE = re.compile(r"\[\[([^\]|]+)\|([^\]]+)\]\]")
WIKILINK_RE = re.compile(r"\[\[([^\]]+)\]\]")
IMAGE_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")


def preprocess_md(text: str) -> str:
    text = FRONTMATTER_RE.sub("", text, count=1)
    text = text.lstrip("\n")
    text = re.sub(r"^# [^\n]+\n", "", text, count=1)
    text = re.sub(r"^(?:>[^\n]*\n)+\s*\n", "", text.lstrip("\n"), count=1)
    text = resolve_transcludes(text)
    text = WIKILINK_ALIAS_RE.sub(r"\2", text)
    text = WIKILINK_RE.sub(r"\1", text)
    return text


def resolve_image(path: str) -> Path | None:
    md_dir = SRC_MD.parent
    p = (md_dir / path).resolve()
    if p.exists():
        return p
    p2 = (ROOT / path.lstrip("/")).resolve()
    if p2.exists():
        return p2
    return None


# ═══════════════════════════════════════════════════════════════════════════
# 5. DOCX rendering primitives
# ═══════════════════════════════════════════════════════════════════════════

HEADING_MAP = {1: "Heading 1", 2: "Heading 1", 3: "Heading 2",
               4: "Heading 3", 5: "Heading 4", 6: "Heading 5"}


def add_runs(paragraph, text: str):
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    for part in pattern.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2]); r.bold = True
        elif part.startswith("*") and part.endswith("*"):
            r = paragraph.add_run(part[1:-1]); r.italic = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1]); r.font.name = "Consolas"
        else:
            paragraph.add_run(part)


def parse_table_rows(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        if re.match(r"^\|[\s:\-|]+\|$", line):
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
    return rows


def _set_cell_shading(cell, hex_color: str):
    from docx.oxml import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd"); tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)


def _set_cell_borders(cell):
    from docx.oxml import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders"); tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "808080")
        tc_borders.append(b)


def add_heading(doc: Document, level: int, text: str):
    p = doc.add_paragraph(style=HEADING_MAP.get(level, "Heading 3"))
    add_runs(p, text)


def add_body_paragraph(doc: Document, text: str):
    try:
        p = doc.add_paragraph(style="normal1")
    except KeyError:
        p = doc.add_paragraph()
    add_runs(p, text)


def add_bullet(doc: Document, text: str, level: int = 0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75 + level * 0.5)
    add_runs(p, f"• {text}")


def add_numbered(doc: Document, text: str, n: int):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    add_runs(p, f"{n}. {text}")


def add_blockquote(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    add_runs(p, text)
    for r in p.runs:
        r.italic = True
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_image_paragraph(doc: Document, alt: str, src_path: Path, width_cm: float = 15.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        run = p.add_run()
        run.add_picture(str(src_path), width=Cm(width_cm))
    except Exception as e:
        p.add_run(f"[Imagen no disponible: {alt}] ({e})").italic = True
        return
    if alt:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run(alt); cr.italic = True; cr.font.size = Pt(9)


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    rows = [r + [""] * (n_cols - len(r)) for r in rows]
    table = doc.add_table(rows=len(rows), cols=n_cols)
    available = {s.name for s in doc.styles}
    for candidate in ("Light Grid Accent 1", "Table Grid", "Grid Table Light"):
        if candidate in available:
            try:
                table.style = candidate; break
            except KeyError:
                continue
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            add_runs(p, cell_text)
            _set_cell_borders(cell)
            if i == 0:
                _set_cell_shading(cell, "1F4E79")
                for r in p.runs:
                    r.bold = True
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


# ═══════════════════════════════════════════════════════════════════════════
# 6. Main block processor
# ═══════════════════════════════════════════════════════════════════════════


def render_md(doc: Document, md_text: str, conn: sqlite3.Connection):
    lines = md_text.splitlines()
    i, n = 0, len(lines)
    skip_toc = False
    mermaid_count = 0
    sqlseal_count = 0

    while i < n:
        line = lines[i]
        stripped = line.strip()

        if re.match(r"^\s*(---+|\*\*\*+)\s*$", line):
            i += 1; continue

        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            j = i + 1
            code_lines: list[str] = []
            while j < n and not lines[j].strip().startswith("```"):
                code_lines.append(lines[j]); j += 1
            code = "\n".join(code_lines)
            if lang.startswith("sqlseal"):
                sqlseal_count += 1
                try:
                    headers, rows = run_sqlseal(conn, code)
                    if not rows:
                        p = doc.add_paragraph()
                        r = p.add_run("[Consulta sin resultados]")
                        r.italic = True; r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
                    else:
                        add_table(doc, [headers, *rows])
                except Exception as e:
                    print(f"  ⚠ SQLSeal error: {e}")
                    p = doc.add_paragraph()
                    r = p.add_run(f"[Error SQLSeal: {e}]")
                    r.italic = True; r.font.color.rgb = RGBColor(0xB0, 0x20, 0x20)
            elif lang == "mermaid":
                mermaid_count += 1
                print(f"  Rendering Mermaid diagram {mermaid_count}…")
                png = render_mermaid_to_png(code)
                if png:
                    add_image_paragraph(doc, "", png, width_cm=15.0)
                else:
                    p = doc.add_paragraph()
                    r = p.add_run("[Diagrama Mermaid no renderizable]")
                    r.italic = True; r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            else:
                for cl in code_lines:
                    p = doc.add_paragraph()
                    r = p.add_run(cl if cl else " ")
                    r.font.name = "Consolas"; r.font.size = Pt(9)
            i = j + 1
            continue

        m = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if m:
            level = len(m.group(1)); text = m.group(2)
            if level == 2 and text.strip() == "Tabla de Contenido":
                skip_toc = True; i += 1; continue
            if skip_toc and level < 3:
                skip_toc = False
            if skip_toc:
                i += 1; continue
            add_heading(doc, level, text)
            i += 1
            continue

        if skip_toc:
            i += 1; continue

        if stripped.startswith("|") and i + 1 < n and re.match(r"^\s*\|[\s:\-|]+\|\s*$", lines[i + 1]):
            table_lines = [line, lines[i + 1]]
            j = i + 2
            while j < n and lines[j].strip().startswith("|"):
                table_lines.append(lines[j]); j += 1
            rows = parse_table_rows(table_lines)
            add_table(doc, rows)
            doc.add_paragraph()
            i = j
            continue

        img_m = IMAGE_RE.match(stripped)
        if img_m:
            alt, src = img_m.group(1), img_m.group(2)
            resolved = resolve_image(src)
            if resolved:
                add_image_paragraph(doc, alt, resolved)
            else:
                p = doc.add_paragraph()
                p.add_run(f"[Imagen no encontrada: {src}]").italic = True
            i += 1
            continue

        if stripped.startswith(">"):
            quote_lines = []
            while i < n and lines[i].strip().startswith(">"):
                quote_lines.append(lines[i].strip().lstrip(">").strip()); i += 1
            add_blockquote(doc, " ".join(quote_lines))
            continue

        if re.match(r"^\s*[-*]\s+", line):
            while i < n and re.match(r"^\s*[-*]\s+", lines[i]):
                bm = re.match(r"^(\s*)[-*]\s+(.+?)\s*$", lines[i])
                indent = len(bm.group(1)) // 2
                add_bullet(doc, bm.group(2), level=indent)
                i += 1
            continue

        if re.match(r"^\s*\d+\.\s+", line):
            while i < n and re.match(r"^\s*\d+\.\s+", lines[i]):
                nm = re.match(r"^\s*(\d+)\.\s+(.+?)\s*$", lines[i])
                add_numbered(doc, nm.group(2), int(nm.group(1)))
                i += 1
            continue

        if not stripped:
            i += 1; continue

        para_lines = [line.rstrip()]
        j = i + 1
        while j < n:
            nxt = lines[j]; nxt_strip = nxt.strip()
            if not nxt_strip: break
            if nxt_strip.startswith(("#", ">", "-", "*", "|", "```")): break
            if re.match(r"^\d+\.\s+", nxt_strip): break
            if IMAGE_RE.match(nxt_strip): break
            para_lines.append(nxt.rstrip())
            j += 1
        add_body_paragraph(doc, " ".join(para_lines))
        i = j

    print(f"  • Mermaid diagrams rendered: {mermaid_count}")
    print(f"  • SQLSeal queries executed: {sqlseal_count}")


# ═══════════════════════════════════════════════════════════════════════════
# 7. Cover + body assembly
# ═══════════════════════════════════════════════════════════════════════════


def fill_cover(doc: Document):
    for p in doc.paragraphs:
        for placeholder, value in COVER_REPLACEMENTS.items():
            if placeholder in p.text:
                full = p.text.replace(placeholder, value)
                for r in list(p.runs):
                    r.text = ""
                if p.runs:
                    p.runs[0].text = full
                else:
                    p.add_run(full)
    for idx, p in enumerate(doc.paragraphs):
        if "Elkin Cerda González" in p.text:
            new_p = copy.deepcopy(p._p)
            p._p.addnext(new_p)
            new_para = doc.paragraphs[idx + 1]
            for r in list(new_para.runs):
                r.text = ""
            if new_para.runs:
                new_para.runs[0].text = "Santiago Martínez Ramírez"
            else:
                new_para.add_run("Santiago Martínez Ramírez")
            break


def cut_body(doc: Document):
    keep_until_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.startswith("Fecha:"):
            keep_until_idx = i; break
    if keep_until_idx is None:
        raise RuntimeError("Could not locate 'Fecha:' marker in template cover")
    for p in doc.paragraphs[keep_until_idx + 1 :]:
        p._element.getparent().remove(p._element)
    for t in list(doc.tables):
        t._element.getparent().remove(t._element)


def main():
    if not TEMPLATE.exists():
        raise FileNotFoundError(TEMPLATE)
    if not SRC_MD.exists():
        raise FileNotFoundError(SRC_MD)

    print("• Building vault SQLite DB …")
    conn = build_vault_db()

    shutil.copy(TEMPLATE, OUT)
    doc = Document(str(OUT))
    print(f"• Template copied → {OUT.relative_to(ROOT)}")
    fill_cover(doc)
    cut_body(doc)

    md_text = SRC_MD.read_text(encoding="utf-8")
    md_text = preprocess_md(md_text)
    print("• Markdown pre-processed (transcludes resolved)")

    render_md(doc, md_text, conn)
    print("• Markdown rendered")

    doc.save(str(OUT))
    size_kb = OUT.stat().st_size / 1024
    print(f"✔ Saved: {OUT.relative_to(ROOT)}  ({size_kb:.1f} KB)")


if __name__ == "__main__":
    main()
